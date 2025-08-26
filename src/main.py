import os
import io
import json
import csv
import re
from pathlib import Path
from typing import List, Dict, Any, Tuple

try:
	from dotenv import load_dotenv  # optional
	load_dotenv()
except Exception:
	pass

import fitz  # PyMuPDF
from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
from docx import Document
from groq import Groq

# -------------------------
# Utilities
# -------------------------

def clean_whitespace(text: str) -> str:
	if not text:
		return ""
	text = text.replace("\x00", " ")
	text = re.sub(r"[\r\t]+", " ", text)
	text = re.sub(r"\s+", " ", text)
	return text.strip()


def fix_hyphenation(text: str) -> str:
	return re.sub(r"([A-Za-z])\-\n([A-Za-z])", r"\1\2", text)


def drop_repeating_headers(pages: List[str]) -> List[str]:
	if not pages:
		return pages
	first_lines = [p.splitlines()[0].strip() for p in pages if p.strip()]
	last_lines = [p.splitlines()[-1].strip() for p in pages if p.strip()]
	head = None
	tail = None
	if first_lines and first_lines.count(first_lines[0]) > len(first_lines) // 2:
		head = first_lines[0]
	if last_lines and last_lines.count(last_lines[0]) > len(last_lines) // 2:
		tail = last_lines[0]
	result = []
	for p in pages:
		lines = p.splitlines()
		if head and lines and lines[0].strip() == head:
			lines = lines[1:]
		if tail and lines and lines[-1].strip() == tail:
			lines = lines[:-1]
		result.append("\n".join(lines))
	return result

# -------------------------
# Extraction
# -------------------------

def _page_text_blocks(page: fitz.Page) -> str:
	try:
		blocks = page.get_text("blocks")
		blocks = sorted(blocks, key=lambda b: (round(b[1], 2), round(b[0], 2)))
		texts = [b[4] for b in blocks if len(b) >= 5 and b[4]]
		return "\n".join(texts)
	except Exception:
		return page.get_text("text")


def _chars_count(text: str) -> int:
	return sum(1 for c in text if not c.isspace())


def _ocr_page_image(img: Image.Image, lang: str = "eng") -> str:
	return pytesseract.image_to_string(img, lang=lang, config="--oem 1 --psm 6")


def extract_pdf_text(data: bytes, ocr_on_demand: bool = True, lang: str = "eng", low_char_threshold: int = 200) -> str:
	doc = fitz.open(stream=data, filetype="pdf")
	per_page_text: List[str] = []
	low_text_pages: List[int] = []
	for i in range(len(doc)):
		page = doc.load_page(i)
		text = _page_text_blocks(page)
		if _chars_count(text) < low_char_threshold:
			low_text_pages.append(i)
		per_page_text.append(text)
	if ocr_on_demand and low_text_pages:
		images: List[Image.Image] = convert_from_bytes(data, fmt="png")
		for idx in low_text_pages:
			ocr_text = _ocr_page_image(images[idx], lang=lang)
			per_page_text[idx] = ocr_text
	per_page_text = drop_repeating_headers(per_page_text)
	joined = "\n\n".join(per_page_text)
	joined = fix_hyphenation(joined)
	joined = clean_whitespace(joined)
	return joined


def extract_docx_text(data: bytes) -> str:
	doc = Document(io.BytesIO(data))
	parts: List[str] = []
	for p in doc.paragraphs:
		if p.text and p.text.strip():
			parts.append(p.text)
	for tbl in doc.tables:
		for row in tbl.rows:
			cells = [cell.text.strip() for cell in row.cells]
			parts.append("\t".join(cells))
	text = "\n".join(parts)
	return clean_whitespace(text)

# -------------------------
# LLM Scoring
# -------------------------

# SYSTEM = (
# 	"You Become the SME as per the JD"
# 	"You evaluate candidates against a job description from their resume text. "
# 	"Follow instructions. Return JSON only. No extra text."
# )
# PROMPT = (
# 	"You are an ATS evaluator. Return STRICT JSON only, matching this schema: "
# 	"You are an SME based on the JD"
# 	"Always try to analyze the JD and the Resume with the keywords.And anlayze how much suitable the candidate will be on the job role."
# 	"Try to make sense of the skills and projects mentioned in the resume whether they are real or just a false information"
# 	"{\"candidate_name\": string|null, \"final_score\": number, \"hard_filter_pass\": boolean, "
# 	"\"skill_coverage\": number|null, \"project_relevance\": number|null, \"role_alignment\": number|null, "
# 	"\"education_fit\": number|null, \"penalties\": array, \"top_reasons\": array, \"risks\": array, \"evidence_snippets\": array, \"explanation\": string|null}. "
# 	"Always include final_score (0-100). Be concise."
# )
# STRICT_PROMPT = (
# 	"Return JSON with ALL required keys exactly as specified; do not add or omit keys. "
# 	"If unsure, set numeric fields to 0 and arrays to []."
# )
# MAX_CHARS = 20000

SYSTEM = (
    "You are an ATS evaluator and Subject Matter Expert (SME) for the given Job Description (JD). "
    "You will evaluate candidates based on their resume text against the JD. "
    "Return JSON only. No extra text."
)

PROMPT = (
    "Analyze the Job Description (JD) and Resume carefully. Follow these steps:\n"
    "1. Understand the JD completely:\n"
    "   - Extract the core business need (what the role helps the client achieve).\n"
    "   - Identify required years of experience, technical requirements, soft skills, and domain.\n"
    "2. Extract the candidate name from the resume if present\n"
    "3. Compare JD vs Resume:\n"
    "   - Skill coverage: match resume skills against JD requirements.\n"
    "   - Project relevance: check if projects align with JD business need.\n"
    "   - Role alignment: does candidate’s experience level fit JD (lead vs junior)?\n"
    "   - Education fit.\n"
    "4. Compute a final score (0-100) considering all above factors.\n"
    "5. Provide explanations:\n"
    "   - **First**, clearly state the required experience vs candidate experience (e.g., 'Required 7+ years, candidate has 5 years - does not meet criteria' or 'Required 7+ years, candidate has 10 years - meets criteria').\n"
    "   - Then explain which skills, projects, or experience influenced the score (both positive and negative).\n"
    "   - If score decreased, explain why (missing skills, domain mismatch, lack of leadership, etc.).\n"
    "6. Identify penalties for major gaps.\n"
    "7. Return STRICT JSON only in this schema:\n"
    "{"
    "\"candidate_name\": string|null, "
    "\"final_score\": number, "
    "\"hard_filter_pass\": boolean, "
    "\"skill_coverage\": number|null, "
    "\"project_relevance\": number|null, "
    "\"role_alignment\": number|null, "
    "\"education_fit\": number|null, "
    "\"penalties\": array, "
    "\"top_reasons\": array, "
    "\"risks\": array, "
    "\"evidence_snippets\": array, "
    "\"explanation\": string"
    "}.\n"
    "Important:\n"
    "- Do not include any text outside the JSON.\n"
    "- If unsure, set numeric fields to 0 and arrays to [].\n"
)

STRICT_PROMPT = (
	"Return JSON with ALL required keys exactly as specified; do not add or omit keys. "
	"If unsure, set numeric fields to 0 and arrays to []."
    "Ensure the explanation clearly states why the candidate scored high or low (mention skills, projects, experience that influenced the score)."
	"Ensure the explanation starts with the experience check (required vs actual) and then the reasoning for score adjustments."
)

MAX_CHARS = 20000


def _truncate(text: str, max_chars: int = MAX_CHARS) -> str:
	return text[:max_chars] if len(text) > max_chars else text

##################### Groq API #####################################

def _client() -> Groq:
	api_key = os.environ.get("GROQ_API_KEY")
	if not api_key:
		raise RuntimeError("GROQ_API_KEY not set")
	return Groq(api_key=api_key)

def chat_json(model: str, system: str, user: str) -> Dict[str, Any]:
	client = _client()
	resp = client.chat.completions.create(
		model=model,
		messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
		temperature=0.0,
		max_tokens=2048,
	)
	content = resp.choices[0].message.content
	if content.strip().startswith("```"):
		content = content.strip().strip("`")
		content = content.split("\n", 1)[-1]
	return json.loads(content)

def score_resume(model: str, jd_text: str, resume_text: str) -> Dict[str, Any]:
	user = f"<JD>\n{jd_text}\n</JD>\n<RESUME>\n{resume_text}\n</RESUME>\n{PROMPT}"
	try:
		obj = chat_json(model=model, system=SYSTEM, user=user)
		if "final_score" in obj:
			return _ensure_schema(obj)
	except Exception:
		pass
	user2 = f"<JD>\n{jd_text}\n</JD>\n<RESUME>\n{resume_text}\n</RESUME>\n{PROMPT}\n{STRICT_PROMPT}"
	try:
		obj2 = chat_json(model=model, system=SYSTEM, user=user2)
		return _ensure_schema(obj2)
	except Exception:
		return _ensure_schema({"final_score": 0})

##################################################################



##################### Ollama #####################################

# import ollama

# def chat_json_ollama(system: str, user: str, model: str = "llama3:8b") -> dict:
#     """
#     Calls Ollama for chat completion and returns parsed JSON.
#     """
#     messages = [
#         {"role": "system", "content": system},
#         {"role": "user", "content": user}
#     ]
#     resp = ollama.chat(model=model, messages=messages)
#     content = resp['message']['content'].strip()
#     print(content)
#     # Clean code-block formatting if present
#     if content.startswith("```"):
#         content = content.strip("`").split("\n", 1)[-1]
    
#     try:
#         return json.loads(content)
#     except json.JSONDecodeError:
#         # fallback if JSON invalid
#         return {"final_score": 0}

# def score_resume(model: str, jd_text: str, resume_text: str) -> Dict[str, Any]:
#     user = f"<JD>\n{jd_text}\n</JD>\n<RESUME>\n{resume_text}\n</RESUME>\n{PROMPT}"
#     print("\n--- Sending request to model ---")
#     print("JD snippet:", jd_text[:200].replace("\n", " "))
#     print("Resume snippet:", resume_text[:200].replace("\n", " "))
    
#     try:
#         obj = chat_json_ollama(model=model, system=SYSTEM, user=user)
#         print("\n--- Raw model output ---")
#         print(obj)
#         if "final_score" in obj:
#             return _ensure_schema(obj)
#     except Exception as e:
#         print("Error in first attempt:", e)
    
#     # Fallback with STRICT_PROMPT
#     user2 = f"<JD>\n{jd_text}\n</JD>\n<RESUME>\n{resume_text}\n</RESUME>\n{PROMPT}\n{STRICT_PROMPT}"
#     try:
#         obj2 = chat_json_ollama(model=model, system=SYSTEM, user=user2)
#         print("\n--- Raw model output (strict) ---")
#         print(obj2)
#         return _ensure_schema(obj2)
#     except Exception as e:
#         print("Error in fallback attempt:", e)
#         return _ensure_schema({"final_score": 0})


##################################################################

def _ensure_list_str(value: Any) -> List[str]:
	if value is None:
		return []
	if isinstance(value, list):
		return [str(v) for v in value]
	return [str(value)]


def _ensure_schema(obj: Dict[str, Any]) -> Dict[str, Any]:
	def num(v, default=0.0):
		try:
			return float(v)
		except Exception:
			return float(default)
	out: Dict[str, Any] = {}
	out["candidate_name"] = obj.get("candidate_name")
	out["final_score"] = max(0.0, min(100.0, num(obj.get("final_score"), 0)))
	out["hard_filter_pass"] = bool(obj.get("hard_filter_pass", True))
	out["skill_coverage"] = num(obj.get("skill_coverage"), 0)
	out["project_relevance"] = num(obj.get("project_relevance"), 0)
	out["role_alignment"] = num(obj.get("role_alignment"), 0)
	out["education_fit"] = num(obj.get("education_fit"), 0)
	raw_pen = obj.get("penalties") or []
	coerced: List[Dict[str, Any]] = []
	for p in raw_pen if isinstance(raw_pen, list) else [raw_pen]:
		if isinstance(p, dict):
			coerced.append(p)
		else:
			coerced.append({"reason": str(p), "points": 0})
	out["penalties"] = coerced
	out["top_reasons"] = _ensure_list_str(obj.get("top_reasons"))
	out["risks"] = _ensure_list_str(obj.get("risks"))
	out["evidence_snippets"] = _ensure_list_str(obj.get("evidence_snippets"))
	out["explanation"] = obj.get("explanation")
	return out



# -------------------------
# Main
# -------------------------

def run():
	project_root = Path(__file__).resolve().parent.parent
	input_dir = project_root / "resumes"
	out_dir = project_root / "standalone_output"
	jd_file = project_root / "jd" / "jd.txt"
	# model = "llama-3.3-70b-versatile"
	model = "llama3:8b"  # Ollama local model
	batch_size = 5
	limit = 5
	lang = "eng"
	low_char_threshold = 200

	print(f"Input: {input_dir}")
	print(f"Output: {out_dir}")
	print(f"JD: {jd_file}")
	out_dir.mkdir(parents=True, exist_ok=True)

	# 1) Extract
	paths: List[Path] = []
	if input_dir.is_dir():
		for p in sorted(input_dir.iterdir()):
			if p.suffix.lower() in {".pdf", ".docx"}:
				paths.append(p)
	paths = paths[:batch_size]

	for p in paths:
		data = p.read_bytes()
		if p.suffix.lower() == ".pdf":
			text = extract_pdf_text(data, ocr_on_demand=True, lang=lang, low_char_threshold=low_char_threshold)
		elif p.suffix.lower() == ".docx":
			text = extract_docx_text(data)
		else:
			continue
		(out_dir / f"{p.stem}.txt").write_text(text, encoding="utf-8")

	# 2) Score
	if not os.environ.get("GROQ_API_KEY"):
		raise RuntimeError("GROQ_API_KEY not set. Put it in .env or set in env.")
	jd_text = _truncate(jd_file.read_text(encoding="utf-8"))
	txts = sorted(out_dir.glob("*.txt"))[:limit]
	rows: List[Dict[str, Any]] = []
	jsonl_path = out_dir / "scores.jsonl"
	csv_path = out_dir / "scores.csv"
	with jsonl_path.open("w", encoding="utf-8") as jf:
		for t in txts:
			resume_text = _truncate(t.read_text(encoding="utf-8", errors="ignore"))
			rec = score_resume(model, jd_text, resume_text)
			rec["file"] = t.name
			jf.write(json.dumps(rec, ensure_ascii=False) + "\n")
			rows.append(rec)
	rows.sort(key=lambda r: r.get("final_score", 0), reverse=True)
	with csv_path.open("w", encoding="utf-8", newline="") as cf:
		writer = csv.DictWriter(cf, fieldnames=["file", "candidate_name", "final_score", "hard_filter_pass", "top_reasons", "risks"])
		writer.writeheader()
		for r in rows:
			row = {"file": r.get("file"),
					"candidate_name": r.get("candidate_name"),
					"final_score": r.get("final_score"),
					"hard_filter_pass": r.get("hard_filter_pass"),
					"top_reasons": " | ".join(r.get("top_reasons", [])),
					"risks": " | ".join(r.get("risks", []))}
			writer.writerow(row)
	print(f"Done. Wrote {jsonl_path} and {csv_path}")


if __name__ == "__main__":
	run()